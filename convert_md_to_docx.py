"""
Markdown to Word (.docx) Converter
Chuyển đổi báo cáo chương 2, 3, 4 từ Markdown sang Word với formatting đầy đủ.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_styles(doc):
    """Setup custom styles for the document."""
    style = doc.styles

    # Normal style
    normal = style['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    # Set font for East Asian text
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Heading 1
    h1 = style['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 2
    h2 = style['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = style['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = style['Heading 4']
    h4.font.name = 'Times New Roman'
    h4.font.size = Pt(13)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # Create Code style
    try:
        code_style = style.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code_style = style['Code Block']
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(30, 30, 30)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.15
    code_style.paragraph_format.left_indent = Cm(1)

    # Create Note style (for blockquotes)
    try:
        note_style = style.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        note_style = style['Note']
    note_style.font.name = 'Times New Roman'
    note_style.font.size = Pt(11)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(80, 80, 80)
    note_style.paragraph_format.left_indent = Cm(1)
    note_style.paragraph_format.space_before = Pt(4)
    note_style.paragraph_format.space_after = Pt(4)

    return doc


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code) to a paragraph."""
    # Process inline formatting
    # Order: code first, then bold+italic, bold, italic
    parts = re.split(r'(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('`') and part.endswith('`') and len(part) > 1:
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 40, 40)
        elif part.startswith('***') and part.endswith('***') and len(part) > 5:
            # Bold + Italic
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**') and len(part) > 3:
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header.strip())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Clean up markdown formatting in cells
            clean_text = cell_text.strip()
            clean_text = re.sub(r'<br\s*/?>', '\n', clean_text)
            add_formatted_text(p, clean_text)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    # Add spacing after table
    doc.add_paragraph('')
    return table


def parse_table(lines, start_idx):
    """Parse a markdown table starting at the given index."""
    headers = []
    rows = []
    i = start_idx

    # Parse header
    if i < len(lines) and '|' in lines[i]:
        parts = [p.strip() for p in lines[i].split('|')]
        headers = [p for p in parts if p and not all(c in '-: ' for c in p)]
        i += 1

    # Skip separator line (|---|---|...)
    if i < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i]):
        i += 1

    # Parse data rows
    while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
        parts = [p.strip() for p in lines[i].split('|')]
        row = [p for p in parts if p != '']
        # Handle case where split creates empty strings at start/end
        if row:
            rows.append(row)
        i += 1

    return headers, rows, i


def is_mermaid_or_code_start(line):
    """Check if line starts a code block."""
    stripped = line.strip()
    return stripped.startswith('```')


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    print(f"  Đang chuyển đổi: {os.path.basename(md_path)} → {os.path.basename(docx_path)}")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks (```)
        if stripped.startswith('```') and not in_code_block:
            in_code_block = True
            code_lang = stripped[3:].strip()
            code_lines = []
            i += 1
            continue
        elif stripped.startswith('```') and in_code_block:
            in_code_block = False
            # Add code block to document
            if code_lang.lower() == 'mermaid':
                # For mermaid diagrams, add as a note
                p = doc.add_paragraph()
                p.style = doc.styles['Note']
                run = p.add_run('[Sơ đồ Mermaid — Xem bản gốc Markdown để hiển thị sơ đồ tương tác]')
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                # Still add the mermaid code for reference
                if code_lines:
                    # Add a light description of the diagram
                    desc_lines = []
                    for cl in code_lines:
                        cl_stripped = cl.strip()
                        if cl_stripped.startswith('subgraph'):
                            name = cl_stripped.replace('subgraph', '').strip().strip('"[]')
                            desc_lines.append(f'  • Nhóm: {name}')
                        elif '-->|' in cl_stripped or '-->' in cl_stripped:
                            # Connection line - extract readable form
                            clean = cl_stripped.replace('-->', ' → ').replace('-->|', ' → (').replace('|', ')')
                            clean = re.sub(r'[\[\]"{}]', '', clean)
                            if clean.strip():
                                desc_lines.append(f'  {clean.strip()}')

                    if desc_lines:
                        for dl in desc_lines[:8]:  # Limit to 8 lines
                            p2 = doc.add_paragraph()
                            p2.style = doc.styles['Note']
                            p2.add_run(dl)
            else:
                # Regular code block
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    # Add language label
                    if code_lang:
                        p_label = doc.add_paragraph()
                        run_label = p_label.add_run(f'[{code_lang.upper()}]')
                        run_label.font.name = 'Consolas'
                        run_label.font.size = Pt(8)
                        run_label.font.color.rgb = RGBColor(100, 100, 100)
                        run_label.bold = True
                        p_label.paragraph_format.space_after = Pt(0)
                        p_label.paragraph_format.space_before = Pt(8)

                    for code_line in code_lines:
                        p = doc.add_paragraph()
                        p.style = doc.styles['Code Block']
                        run = p.add_run(code_line)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)

            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle horizontal rules (---)
        if re.match(r'^---+\s*$', stripped):
            # Add a thin line / page break hint
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a horizontal line via border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Handle tables
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            # Check if next line is separator
            if i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i + 1]):
                headers, rows, new_i = parse_table(lines, i)
                if headers:
                    add_table(doc, headers, rows)
                i = new_i
                continue

        # Handle headings
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if match:
                level = len(match.group(1))
                heading_text = match.group(2).strip()
                # Clean markdown formatting from heading
                heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
                heading_text = re.sub(r'\*(.+?)\*', r'\1', heading_text)
                heading_text = re.sub(r'`(.+?)`', r'\1', heading_text)

                heading_style = f'Heading {min(level, 4)}'
                p = doc.add_heading(heading_text, level=min(level, 4))
                # Ensure Times New Roman for headings
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                i += 1
                continue

        # Handle blockquotes
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            # Remove alert markers
            quote_text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', quote_text).strip()
            if quote_text:
                p = doc.add_paragraph()
                p.style = doc.styles['Note']
                add_formatted_text(p, quote_text)
            i += 1
            continue

        # Handle unordered lists
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', stripped)
        if list_match:
            indent_level = len(line) - len(line.lstrip())
            list_text = list_match.group(2).strip()

            p = doc.add_paragraph()
            p.style = 'List Bullet'
            if indent_level >= 4:
                p.paragraph_format.left_indent = Cm(2.5)
            elif indent_level >= 2:
                p.paragraph_format.left_indent = Cm(1.5)

            add_formatted_text(p, list_text)
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
            i += 1
            continue

        # Handle ordered lists
        ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', stripped)
        if ordered_match:
            list_text = ordered_match.group(3).strip()
            num = ordered_match.group(2)

            p = doc.add_paragraph()
            p.style = 'List Number'
            add_formatted_text(p, list_text)
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
            i += 1
            continue

        # Handle empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        # Handle indented text
        indent = len(line) - len(line.lstrip())
        if indent >= 4:
            p.paragraph_format.left_indent = Cm(1)
        add_formatted_text(p, stripped)
        for run in p.runs:
            if run.font.name is None:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        p.paragraph_format.first_line_indent = Cm(1)
        i += 1

    doc.save(docx_path)
    print(f"  ✅ Đã lưu: {docx_path}")
    return docx_path


def main():
    print("=" * 60)
    print("  CHUYỂN ĐỔI BÁO CÁO MARKDOWN → WORD (.DOCX)")
    print("=" * 60)

    # Define input/output paths
    chapters = [
        {
            'md': r'c:\Users\Admin\.gemini\antigravity-ide\brain\fdc478b8-6de4-4e65-bdfa-adbe81d8a24b\bao_cao_chuong_2.md',
            'docx': r'c:\tmp\autograder\Bao_cao_Chuong_2.docx',
            'name': 'Chương 2: Phân tích và Thiết kế Hệ thống'
        },
        {
            'md': r'c:\Users\Admin\.gemini\antigravity-ide\brain\5d54746a-1647-411c-808c-f2719b2a73e3\bao_cao_chuong_3.md',
            'docx': r'c:\tmp\autograder\Bao_cao_Chuong_3.docx',
            'name': 'Chương 3: Hiện thực hóa Hệ thống'
        },
        {
            'md': r'c:\Users\Admin\.gemini\antigravity-ide\brain\5d54746a-1647-411c-808c-f2719b2a73e3\bao_cao_chuong_4.md',
            'docx': r'c:\tmp\autograder\Bao_cao_Chuong_4.docx',
            'name': 'Chương 4: Kiểm thử, Đánh giá và Kết luận'
        }
    ]

    results = []
    for ch in chapters:
        print(f"\n📄 {ch['name']}")
        if not os.path.exists(ch['md']):
            print(f"  ❌ Không tìm thấy file: {ch['md']}")
            continue
        try:
            path = convert_md_to_docx(ch['md'], ch['docx'])
            results.append(path)
        except Exception as e:
            print(f"  ❌ Lỗi: {e}")
            import traceback
            traceback.print_exc()

    # Also create a combined document
    print(f"\n📚 Tạo file tổng hợp tất cả chương...")
    try:
        combined_doc = Document()
        setup_styles(combined_doc)
        for section in combined_doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)

        for idx, ch in enumerate(chapters):
            if not os.path.exists(ch['md']):
                continue

            with open(ch['md'], 'r', encoding='utf-8') as f:
                content = f.read()

            if idx > 0:
                combined_doc.add_page_break()

            # Reuse the same parsing logic
            lines = content.split('\n')
            i = 0
            in_code_block = False
            code_lines = []
            code_lang = ''

            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                if stripped.startswith('```') and not in_code_block:
                    in_code_block = True
                    code_lang = stripped[3:].strip()
                    code_lines = []
                    i += 1
                    continue
                elif stripped.startswith('```') and in_code_block:
                    in_code_block = False
                    if code_lang.lower() == 'mermaid':
                        p = combined_doc.add_paragraph()
                        p.style = combined_doc.styles['Note']
                        run = p.add_run('[Sơ đồ Mermaid — Xem bản gốc Markdown để hiển thị sơ đồ tương tác]')
                        run.italic = True
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    else:
                        code_text = '\n'.join(code_lines)
                        if code_text.strip():
                            if code_lang:
                                p_label = combined_doc.add_paragraph()
                                run_label = p_label.add_run(f'[{code_lang.upper()}]')
                                run_label.font.name = 'Consolas'
                                run_label.font.size = Pt(8)
                                run_label.font.color.rgb = RGBColor(100, 100, 100)
                                run_label.bold = True
                                p_label.paragraph_format.space_after = Pt(0)
                                p_label.paragraph_format.space_before = Pt(8)
                            for code_line in code_lines:
                                p = combined_doc.add_paragraph()
                                p.style = combined_doc.styles['Code Block']
                                run = p.add_run(code_line)
                                run.font.name = 'Consolas'
                                run.font.size = Pt(9)
                    i += 1
                    continue

                if in_code_block:
                    code_lines.append(line)
                    i += 1
                    continue

                if re.match(r'^---+\s*$', stripped):
                    p = combined_doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(12)
                    pPr = p._p.get_or_add_pPr()
                    pBdr = parse_xml(
                        f'<w:pBdr {nsdecls("w")}>'
                        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                        f'</w:pBdr>'
                    )
                    pPr.append(pBdr)
                    i += 1
                    continue

                if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
                    if i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|', lines[i + 1]):
                        headers, rows, new_i = parse_table(lines, i)
                        if headers:
                            add_table(combined_doc, headers, rows)
                        i = new_i
                        continue

                if stripped.startswith('#'):
                    match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
                    if match:
                        level = len(match.group(1))
                        heading_text = match.group(2).strip()
                        heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
                        heading_text = re.sub(r'\*(.+?)\*', r'\1', heading_text)
                        heading_text = re.sub(r'`(.+?)`', r'\1', heading_text)
                        p = combined_doc.add_heading(heading_text, level=min(level, 4))
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                        i += 1
                        continue

                if stripped.startswith('>'):
                    quote_text = stripped.lstrip('>').strip()
                    quote_text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', quote_text).strip()
                    if quote_text:
                        p = combined_doc.add_paragraph()
                        p.style = combined_doc.styles['Note']
                        add_formatted_text(p, quote_text)
                    i += 1
                    continue

                list_match = re.match(r'^(\s*)[-*]\s+(.+)$', stripped)
                if list_match:
                    indent_level = len(line) - len(line.lstrip())
                    list_text = list_match.group(2).strip()
                    p = combined_doc.add_paragraph()
                    p.style = 'List Bullet'
                    if indent_level >= 4:
                        p.paragraph_format.left_indent = Cm(2.5)
                    elif indent_level >= 2:
                        p.paragraph_format.left_indent = Cm(1.5)
                    add_formatted_text(p, list_text)
                    for run in p.runs:
                        if run.font.name is None:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                    i += 1
                    continue

                ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', stripped)
                if ordered_match:
                    list_text = ordered_match.group(3).strip()
                    p = combined_doc.add_paragraph()
                    p.style = 'List Number'
                    add_formatted_text(p, list_text)
                    for run in p.runs:
                        if run.font.name is None:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                    i += 1
                    continue

                if not stripped:
                    i += 1
                    continue

                p = combined_doc.add_paragraph()
                indent = len(line) - len(line.lstrip())
                if indent >= 4:
                    p.paragraph_format.left_indent = Cm(1)
                add_formatted_text(p, stripped)
                for run in p.runs:
                    if run.font.name is None:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                p.paragraph_format.first_line_indent = Cm(1)
                i += 1

        combined_path = r'c:\tmp\autograder\Bao_cao_Day_du_Chuong_2_3_4.docx'
        combined_doc.save(combined_path)
        print(f"  ✅ Đã lưu file tổng hợp: {combined_path}")
        results.append(combined_path)
    except Exception as e:
        print(f"  ❌ Lỗi tạo file tổng hợp: {e}")
        import traceback
        traceback.print_exc()

    print("\n" + "=" * 60)
    print(f"  HOÀN TẤT: {len(results)} file đã được tạo thành công!")
    print("=" * 60)
    for r in results:
        size_kb = os.path.getsize(r) / 1024
        print(f"  📄 {os.path.basename(r)} ({size_kb:.0f} KB)")
    print()


if __name__ == '__main__':
    main()
