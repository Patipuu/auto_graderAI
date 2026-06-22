import io
from flask import Blueprint, request, jsonify
from services.question_service import QuestionService

question_bp = Blueprint('question_bp', __name__)

@question_bp.route('/api/questions', methods=['GET'])
def get_questions():
    """Fetch all questions from bank"""
    query = request.args.get('q', '').lower()
    subject = request.args.get('subject', '')
    
    questions = QuestionService.get_questions(query, subject)
    return jsonify(questions), 200

@question_bp.route('/api/questions/import', methods=['POST'])
def import_questions():
    """Import bulk questions from parsed text"""
    data = request.json or {}
    questions = data.get('questions', [])

    imported, error = QuestionService.persist_imported_questions(questions)
    if error:
        payload, status = error
        return jsonify(payload), status
    return jsonify({'success': True, 'count': len(imported), 'questions': imported}), 201

@question_bp.route('/api/questions/import-text', methods=['POST'])
def import_questions_from_text():
    """Parse and import questions from raw text on the backend."""
    from services.ai_service import parse_questions_from_text

    data = request.json or {}
    text = data.get('text', '')
    subject = data.get('subject', 'Chung')
    difficulty = data.get('difficulty', 'Trung bình')

    questions = parse_questions_from_text(text, subject=subject, difficulty=difficulty)
    imported, error = QuestionService.persist_imported_questions(questions)
    if error:
        payload, status = error
        return jsonify(payload), status
    return jsonify({'success': True, 'count': len(imported), 'questions': imported}), 201

@question_bp.route('/api/questions/import-file', methods=['POST'])
def import_questions_from_file():
    """Extract text from txt/docx/pdf and import parsed questions."""
    from services.ai_service import parse_questions_from_text

    upload = request.files.get('file')
    if not upload:
        return jsonify({'success': False, 'message': 'Missing file'}), 400

    subject = request.form.get('subject', 'Chung')
    difficulty = request.form.get('difficulty', 'Trung bình')
    filename = (upload.filename or '').lower()
    raw = upload.read()

    try:
        if filename.endswith('.txt'):
            text = raw.decode('utf-8-sig')
        elif filename.endswith('.docx'):
            from docx import Document
            document = Document(io.BytesIO(raw))
            text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
        elif filename.endswith('.pdf'):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = '\n'.join((page.extract_text() or '') for page in reader.pages)
        else:
            return jsonify({'success': False, 'message': 'Only .txt, .docx, and .pdf files are supported'}), 400
    except Exception as e:
        return jsonify({'success': False, 'message': f'Cannot extract text from file: {str(e)}'}), 400

    questions = parse_questions_from_text(text, subject=subject, difficulty=difficulty)
    imported, error = QuestionService.persist_imported_questions(questions)
    if error:
        payload, status = error
        return jsonify(payload), status
    return jsonify({'success': True, 'count': len(imported), 'questions': imported}), 201

@question_bp.route('/api/questions/parse-preview', methods=['POST'])
def parse_questions_preview():
    """Parse questions from text without persisting them."""
    from services.ai_service import parse_questions_from_text

    data = request.json or {}
    text = data.get('text', '')
    subject = data.get('subject', 'Chung')
    difficulty = data.get('difficulty', 'Trung bình')

    questions = parse_questions_from_text(text, subject=subject, difficulty=difficulty)
    return jsonify({'success': True, 'questions': questions}), 200

@question_bp.route('/api/questions/parse-preview-file', methods=['POST'])
def parse_questions_preview_file():
    """Extract text from txt/docx/pdf and parse questions without persisting."""
    from services.ai_service import parse_questions_from_text

    upload = request.files.get('file')
    if not upload:
        return jsonify({'success': False, 'message': 'Missing file'}), 400

    subject = request.form.get('subject', 'Chung')
    difficulty = request.form.get('difficulty', 'Trung bình')
    filename = (upload.filename or '').lower()
    raw = upload.read()

    try:
        if filename.endswith('.txt'):
            text = raw.decode('utf-8-sig')
        elif filename.endswith('.docx'):
            from docx import Document
            document = Document(io.BytesIO(raw))
            text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
        elif filename.endswith('.pdf'):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = '\n'.join((page.extract_text() or '') for page in reader.pages)
        else:
            return jsonify({'success': False, 'message': 'Only .txt, .docx, and .pdf files are supported'}), 400
    except Exception as e:
        return jsonify({'success': False, 'message': f'Cannot extract text from file: {str(e)}'}), 400

    questions = parse_questions_from_text(text, subject=subject, difficulty=difficulty)
    return jsonify({'success': True, 'questions': questions}), 200

@question_bp.route('/api/questions/<question_id>', methods=['DELETE'])
def delete_question(question_id):
    """Delete a question from bank"""
    success, error = QuestionService.delete_question(question_id)
    if error:
        return jsonify({'success': False, 'message': error}), 404
    return jsonify({'success': True}), 200
